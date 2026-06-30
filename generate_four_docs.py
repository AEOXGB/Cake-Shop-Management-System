
# -*- coding: utf-8 -*-
"""
蛋糕店管理系统 - 更多Word文档生成脚本
生成自动化测试报告、骑手端需求文档、管理员需求文档、项目上线操作文档的Word版本
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re


def set_chinese_font(run, font_name='宋体', size=12, bold=False, color=None):
    """设置中文字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading_styled(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_chinese_font(run, font_name='黑体', size=16 if level == 1 else 14, bold=True)
    return heading


def add_paragraph_styled(doc, text, font_name='宋体', size=12, bold=False, indent=False):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    set_chinese_font(run, font_name=font_name, size=size, bold=bold)
    return p


def add_table_from_data(doc, headers, rows, first_col_bold=False):
    """从数据添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_chinese_font(run, font_name='黑体', size=11, bold=True)

    # 设置数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            is_bold = first_col_bold and col_idx == 0
            set_chinese_font(run, font_name='宋体', size=10.5, bold=is_bold)

    return table


def add_code_block(doc, code_text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(20)
    run = p.add_run(code_text)
    set_chinese_font(run, font_name='Consolas', size=10)
    return p


def add_list_item(doc, text, level=0):
    """添加列表项"""
    p = doc.add_paragraph(style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Pt(20 * level)
    run = p.add_run(text)
    set_chinese_font(run, font_name='宋体', size=12)
    return p


def generate_automation_test_report(output_path):
    """生成自动化测试报告Word文档"""
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CakeShop 蛋糕系统')
    set_chinese_font(run, font_name='黑体', size=32, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('自动化测试报告')
    set_chinese_font(run, font_name='黑体', size=28, bold=True, color=(0, 102, 204))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        '测试日期：2026年6月24日',
        '测试版本：v1.0.0',
        '测试环境：Windows + Spring Boot 2.2.6 + Java 8 + MySQL',
        '测试工具：JUnit 5 + Mockito + Playwright',
    ]

    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_chinese_font(run, font_name='宋体', size=14)

    doc.add_page_break()

    # ===== 目录 =====
    add_heading_styled(doc, '目录', level=1)
    toc_items = [
        '一、测试概述',
        '二、详细测试结果',
        '三、测试覆盖总结',
        '四、PRD需求达成率',
        '五、测试文件清单',
        '六、测试风险与改进建议',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_chinese_font(run, font_name='宋体', size=12)

    doc.add_page_break()

    # ===== 一、测试概述 =====
    add_heading_styled(doc, '一、测试概述', level=1)

    add_heading_styled(doc, '1.1 测试范围', level=2)
    add_paragraph_styled(doc, '本次测试覆盖 CakeShop 蛋糕系统的多个层面，包括功能完整性测试、功能闭环测试、E2E测试、接口测试、单元测试、集成测试、验收测试和响应式UI测试，共计39个测试用例。', indent=True)

    scope_headers = ['测试类型', '说明', '用例数量']
    scope_rows = [
        ['功能完整性测试', '验证每个功能模块是否正常工作', '12'],
        ['功能闭环测试', '验证业务流程的完整闭环', '2'],
        ['E2E测试', '端到端用户流程测试', '4'],
        ['接口测试', 'API接口功能测试', '5'],
        ['单元测试', '各组件的独立测试', '5'],
        ['集成测试', '组件间集成测试', '3'],
        ['验收测试', 'PRD需求完成度检查', '4'],
        ['响应式UI测试', '页面布局和资源加载', '4'],
        ['总计', '', '39'],
    ]
    add_table_from_data(doc, scope_headers, scope_rows, first_col_bold=True)

    add_heading_styled(doc, '1.2 测试工具', level=2)
    tools_items = [
        '后端测试：JUnit 5 + MockMvc + Mockito',
        '前端E2E测试：Playwright (Chromium)',
        '测试报告：Maven Surefire Report + Playwright HTML Report',
    ]
    for item in tools_items:
        add_list_item(doc, item)

    doc.add_page_break()

    # ===== 二、详细测试结果 =====
    add_heading_styled(doc, '二、详细测试结果', level=1)

    add_heading_styled(doc, '2.1 功能完整性测试 (12项)', level=2)
    fa_headers = ['编号', '测试名称', '结果', '说明']
    fa_rows = [
        ['TC-FA-001', '页面可访问性', '✅ 通过', '首页/登录/注册 均可正常访问'],
        ['TC-FA-002', '登录保护', '✅ 通过', '未登录访问受保护页面自动跳转'],
        ['TC-FA-003', '首页加载性能', '✅ 通过', '加载时间 < 5秒'],
        ['TC-FA-004', '导航菜单', '✅ 通过', '首页包含导航菜单'],
        ['TC-FA-005', '链接完整性', '✅ 通过', '首页链接有效'],
        ['TC-FA-006', '商品展示区', '✅ 通过', '首页包含商品展示区域'],
        ['TC-UF-001', '注册页面显示', '✅ 通过', '注册表单元素完整'],
        ['TC-UF-002', '空表单验证', '✅ 通过', '空表单提交有响应'],
        ['TC-UF-003', '登录页面显示', '✅ 通过', '登录表单元素完整'],
        ['TC-UF-004', '错误密码验证', '✅ 通过', '错误密码登录被阻止'],
        ['TC-UF-005', '商品列表显示', '✅ 通过', '商品列表正常显示'],
        ['TC-UF-006', '商品详情页', '✅ 通过', '商品详情页可打开'],
    ]
    add_table_from_data(doc, fa_headers, fa_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.2 功能闭环测试 (2项)', level=2)
    fc_headers = ['编号', '测试名称', '结果', '说明']
    fc_rows = [
        ['TC-UF-007', '购物车访问', '✅ 通过', '已登录用户可访问购物车'],
        ['TC-UF-008', '我的订单访问', '✅ 通过', '已登录用户可查看订单'],
        ['TC-UF-009', '搜索功能', '✅ 通过', '搜索功能可用'],
    ]
    add_table_from_data(doc, fc_headers, fc_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.3 E2E端到端测试 (4项)', level=2)
    e2e_headers = ['编号', '测试名称', '结果', '说明']
    e2e_rows = [
        ['TC-E2E-001', '用户完整购物流程', '✅ 通过', '登录→浏览→购物车→订单'],
        ['TC-E2E-002', '注册到浏览闭环', '✅ 通过', '注册→登录→浏览首页'],
        ['TC-E2E-003', '管理员管理闭环', '✅ 通过', '登录→商品→订单→用户管理'],
        ['TC-E2E-004', '骑手端页面', '✅ 通过', '骑手待接单页面正常'],
    ]
    add_table_from_data(doc, e2e_headers, e2e_rows)

    doc.add_page_break()

    add_heading_styled(doc, '2.4 接口测试 (5项)', level=2)
    api_headers = ['编号', '测试名称', '结果', '说明']
    api_rows = [
        ['TC-API-001', '登录接口', '✅ 通过', 'POST /user/login'],
        ['TC-API-002', '注册接口', '✅ 通过', 'POST /user/register'],
        ['TC-API-003', '商品列表接口', '✅ 通过', 'GET /goods/goodsList'],
        ['TC-API-004', '商品详情接口', '✅ 通过', 'GET /goodsDetail'],
        ['TC-API-005', '创建订单接口', '✅ 通过', 'POST /order/createOrder'],
    ]
    add_table_from_data(doc, api_headers, api_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.5 单元测试 (5项)', level=2)
    ut_headers = ['编号', '测试名称', '结果', '说明']
    ut_rows = [
        ['TC-UT-001', 'UserService注册', '✅ 通过', '用户数据设置验证'],
        ['TC-UT-002', 'UserService查询', '✅ 通过', '用户数据库查询'],
        ['TC-UT-003', 'GoodsService查询', '✅ 通过', '商品查询'],
        ['TC-UT-004', 'GoodsService统计', '✅ 通过', '商品数量统计'],
        ['TC-UT-005', 'OrderService状态流转', '✅ 通过', '订单状态变化验证'],
    ]
    add_table_from_data(doc, ut_headers, ut_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.6 集成测试 (3项)', level=2)
    it_headers = ['编号', '测试名称', '结果', '说明']
    it_rows = [
        ['TC-IT-001', '用户模块数据库集成', '✅ 通过', '注册→数据库写入→查询验证'],
        ['TC-IT-002', '商品与分类关联', '✅ 通过', '商品可关联到分类'],
        ['TC-IT-003', '订单与订单项关联', '✅ 通过', '订单项可关联到订单'],
    ]
    add_table_from_data(doc, it_headers, it_rows)

    doc.add_page_break()

    add_heading_styled(doc, '2.7 验收测试 (4项)', level=2)
    at_headers = ['编号', '测试名称', '结果', '说明']
    at_rows = [
        ['TC-AT-001', '用户端功能清单', '✅ 通过', '首页/列表/新品/热门'],
        ['TC-AT-002', '管理员端功能清单', '✅ 通过', '用户/商品/订单/统计'],
        ['TC-AT-003', '骑手端功能清单', '✅ 通过', '待接单/配送中/已完成'],
        ['TC-AT-004', '订单状态流转验证', '✅ 通过', '状态值定义正确'],
    ]
    add_table_from_data(doc, at_headers, at_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.8 响应式UI测试 (4项)', level=2)
    rl_headers = ['编号', '测试名称', '结果', '说明']
    rl_rows = [
        ['TC-RL-001', '桌面端显示', '✅ 通过', '1920x1080 正常'],
        ['TC-RL-002', '笔记本显示', '✅ 通过', '1366x768 正常'],
        ['TC-RL-003', '平板显示', '✅ 通过', '768x1024 正常'],
        ['TC-RL-004', '移动端显示', '✅ 通过', '375x812 正常'],
    ]
    add_table_from_data(doc, rl_headers, rl_rows)

    doc.add_page_break()

    # ===== 三、测试覆盖总结 =====
    add_heading_styled(doc, '三、测试覆盖总结', level=1)

    add_heading_styled(doc, '3.1 功能模块覆盖', level=2)
    cov_headers = ['模块', '页面/功能', '测试状态']
    cov_rows = [
        ['用户模块', '注册、登录、个人中心', '✅ 已覆盖'],
        ['商品模块', '商品列表、商品详情、搜索', '✅ 已覆盖'],
        ['购物车模块', '添加商品、查看购物车', '✅ 已覆盖'],
        ['订单模块', '创建订单、查看订单', '✅ 已覆盖'],
        ['管理后台', '用户管理、商品管理、订单管理、骑手管理、销售统计、库存预警', '✅ 已覆盖'],
        ['骑手端', '待接单、配送中、已完成', '✅ 已覆盖'],
    ]
    add_table_from_data(doc, cov_headers, cov_rows, first_col_bold=True)

    doc.add_paragraph()

    add_heading_styled(doc, '3.2 测试覆盖率统计', level=2)
    add_paragraph_styled(doc, '测试用例总数: 39，通过: 39，失败: 0，通过率: 100%', indent=True)
    doc.add_paragraph()

    rate_headers = ['测试类型', '覆盖率']
    rate_rows = [
        ['功能完整性测试', '100%'],
        ['功能闭环测试', '100%'],
        ['E2E端到端测试', '100%'],
        ['接口测试', '100%'],
        ['单元测试', '100%'],
        ['集成测试', '100%'],
        ['验收测试', '100%'],
        ['响应式UI测试', '100%'],
    ]
    add_table_from_data(doc, rate_headers, rate_rows, first_col_bold=True)

    doc.add_page_break()

    # ===== 四、PRD需求达成率 =====
    add_heading_styled(doc, '四、PRD需求达成率', level=1)

    add_heading_styled(doc, '4.1 用户端功能达成率: 100%', level=2)
    user_prd_headers = ['需求', '状态', '说明']
    user_prd_rows = [
        ['商品浏览（首页/列表/新品/热门）', '✅', '全部实现'],
        ['商品详情查看', '✅', '包含图片/价格/描述'],
        ['用户注册/登录', '✅', '支持用户名/手机号登录'],
        ['购物车管理', '✅', '添加/查看商品'],
        ['订单管理', '✅', '创建/查看订单'],
        ['热销排行榜', '✅', '页面可访问'],
        ['新品上市', '✅', '页面可访问'],
    ]
    add_table_from_data(doc, user_prd_headers, user_prd_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '4.2 管理员端功能达成率: 100%', level=2)
    admin_prd_headers = ['需求', '状态', '说明']
    admin_prd_rows = [
        ['管理员登录', '✅', '独立登录页面'],
        ['用户管理', '✅', '列表/搜索/审核/冻结'],
        ['商品管理', '✅', '列表/搜索/添加/编辑/删除'],
        ['订单管理', '✅', '列表/详情/发货'],
        ['分类管理', '✅', '列表/添加/编辑/删除'],
        ['骑手管理', '✅', '列表/审核/查看'],
        ['销售统计', '✅', '数据统计页面'],
        ['库存预警', '✅', '低库存商品展示'],
        ['系统设置', '✅', '页面可访问'],
        ['佣金设置', '✅', '页面可访问'],
    ]
    add_table_from_data(doc, admin_prd_headers, admin_prd_rows)

    doc.add_page_break()

    add_heading_styled(doc, '4.3 骑手端功能达成率: 100%', level=2)
    rider_prd_headers = ['需求', '状态', '说明']
    rider_prd_rows = [
        ['骑手登录', '✅', '共用登录页面自动识别'],
        ['待接单列表', '✅', '状态=3待配送订单'],
        ['配送中列表', '✅', '已接单配送中订单'],
        ['已完成列表', '✅', '已完成配送订单'],
        ['个人中心', '✅', '个人信息展示'],
        ['收入明细', '✅', '配送收入统计'],
        ['订单详情', '✅', '可查看配送订单详情'],
    ]
    add_table_from_data(doc, rider_prd_headers, rider_prd_rows)

    doc.add_page_break()

    # ===== 五、测试文件清单 =====
    add_heading_styled(doc, '五、测试文件清单', level=1)

    add_heading_styled(doc, '5.1 后端测试文件', level=2)
    backend_test_headers = ['文件', '说明']
    backend_test_rows = [
        ['ComprehensiveTestSuite.java', '综合测试套件（33个测试）'],
        ['UserServiceTest.java', '用户Service单元测试'],
        ['GoodsServiceTest.java', '商品Service单元测试'],
        ['OrderServiceTest.java', '订单Service单元测试'],
        ['RiderServiceTest.java', '骑手Service单元测试'],
        ['UserControllerTest.java', '用户Controller测试'],
    ]
    add_table_from_data(doc, backend_test_headers, backend_test_rows, first_col_bold=True)

    doc.add_paragraph()

    add_heading_styled(doc, '5.2 前端E2E测试文件', level=2)
    e2e_test_headers = ['文件', '说明']
    e2e_test_rows = [
        ['01-page-access.spec.js', '页面可访问性测试'],
        ['02-user-flow.spec.js', '用户流程测试'],
        ['03-admin-flow.spec.js', '管理员流程测试'],
        ['04-e2e-comprehensive.spec.js', '端到端综合测试'],
        ['05-responsive-ui.spec.js', '响应式UI测试'],
        ['helpers.js', '测试辅助函数'],
    ]
    add_table_from_data(doc, e2e_test_headers, e2e_test_rows, first_col_bold=True)

    doc.add_page_break()

    # ===== 六、测试风险与改进建议 =====
    add_heading_styled(doc, '六、测试风险与改进建议', level=1)

    add_heading_styled(doc, '6.1 已知风险', level=2)
    risk_items = [
        '数据库依赖：现有测试依赖MySQL数据库，建议引入H2内存数据库进行单元测试',
        '会话管理：部分控制器路径重定向（302）需跟踪，建议统一路由设计',
        '测试数据：测试间数据存在耦合，建议使用 @BeforeEach 清理测试数据',
    ]
    for item in risk_items:
        add_list_item(doc, item)

    add_heading_styled(doc, '6.2 改进建议', level=2)
    sug_items = [
        '添加性能测试：使用 JMeter 或 Gatling 进行压力测试',
        '添加安全测试：SQL注入、XSS攻击测试',
        '添加CI/CD集成：配置 GitHub Actions 自动运行测试',
        '提升测试覆盖率：将前端E2E测试集成到 Maven 构建流程',
    ]
    for item in sug_items:
        add_list_item(doc, item)

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('报告生成时间：2026年6月24日  |  生成工具：Playwright + Maven Surefire')
    set_chinese_font(run, font_name='宋体', size=10, color=(128, 128, 128))

    # 保存文档
    doc.save(output_path)
    print(f'自动化测试报告已生成：{output_path}')


def generate_rider_requirements(output_path):
    """生成骑手端需求文档Word文档"""
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('蛋糕系统')
    set_chinese_font(run, font_name='黑体', size=36, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('骑手端功能需求文档')
    set_chinese_font(run, font_name='黑体', size=28, bold=True, color=(204, 102, 0))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_headers = ['项目', '内容']
    info_rows = [
        ['文档版本', 'v1.0'],
        ['创建日期', '2026-06-09'],
        ['适用系统', '蛋糕系统 - 骑手端'],
    ]
    add_table_from_data(doc, info_headers, info_rows, first_col_bold=True)

    doc.add_page_break()

    # ===== 目录 =====
    add_heading_styled(doc, '目录', level=1)
    toc_items = [
        '1. 项目概述',
        '2. 功能需求清单',
        '3. 订单配送模块',
        '4. 个人中心模块',
        '5. 技术架构',
        '6. 页面原型说明',
        '附录：订单状态流转',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_chinese_font(run, font_name='宋体', size=12)

    doc.add_page_break()

    # ===== 1. 项目概述 =====
    add_heading_styled(doc, '1. 项目概述', level=1)

    add_heading_styled(doc, '1.1 项目背景', level=2)
    add_paragraph_styled(doc, '骑手端是蛋糕系统的重要组成部分，为配送骑手提供订单接收、配送管理、路线导航等核心功能，实现订单配送全流程管理。', indent=True)

    add_heading_styled(doc, '1.2 目标用户', level=2)
    add_list_item(doc, '配送骑手：负责订单的配送工作')

    add_heading_styled(doc, '1.3 核心价值', level=2)
    value_items = [
        '高效接收和处理配送订单',
        '实时更新配送状态',
        '提供导航路线指引',
        '统计配送业绩数据',
    ]
    for item in value_items:
        add_list_item(doc, item)

    doc.add_page_break()

    # ===== 2. 功能需求清单 =====
    add_heading_styled(doc, '2. 功能需求清单', level=1)

    req_headers = ['序号', '模块', '功能点', '需求描述', '优先级']
    req_rows = [
        ['1', '订单配送', '订单列表', '展示待接单、待配送、已完成订单', '高'],
        ['2', '订单配送', '接单功能', '骑手抢单或自动分配订单', '高'],
        ['3', '订单配送', '订单详情', '查看订单完整信息（收货信息、商品明细）', '高'],
        ['4', '订单配送', '取货确认', '确认已到店取货', '高'],
        ['5', '订单配送', '开始配送', '更新订单状态为配送中', '高'],
        ['6', '订单配送', '送达确认', '确认订单已送达', '高'],
        ['7', '订单配送', '导航功能', '提供配送路线导航', '高'],
        ['8', '订单配送', '异常处理', '处理配送异常（地址错误、联系不上等）', '中'],
        ['9', '个人中心', '个人信息', '查看和修改个人资料', '中'],
        ['10', '个人中心', '配送统计', '统计配送数量、完成率、评分', '中'],
        ['11', '个人中心', '收入明细', '查看配送收入记录', '中'],
        ['12', '个人中心', '消息通知', '接收系统通知和订单消息', '低'],
    ]
    add_table_from_data(doc, req_headers, req_rows)

    doc.add_page_break()

    # ===== 3. 订单配送模块 =====
    add_heading_styled(doc, '3. 订单配送模块', level=1)

    add_heading_styled(doc, '3.1 订单列表', level=2)
    add_paragraph_styled(doc, '功能描述：展示骑手相关的所有订单，按状态分类', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '页面结构：', bold=True)
    add_list_item(doc, '底部Tab切换：待接单、待配送、已完成')
    add_list_item(doc, '订单卡片：订单号、商品数量、配送地址、预计送达时间、操作按钮')

    doc.add_paragraph()
    add_paragraph_styled(doc, '订单卡片信息：', bold=True)
    card_headers = ['信息项', '说明']
    card_rows = [
        ['订单号', '唯一订单标识'],
        ['商品数量', '订单商品总数'],
        ['配送地址', '收货地址简要信息'],
        ['预计送达时间', '系统计算的预计时间'],
        ['订单金额', '订单总额（用于参考）'],
    ]
    add_table_from_data(doc, card_headers, card_rows, first_col_bold=True)

    doc.add_paragraph()

    add_heading_styled(doc, '3.2 接单功能', level=2)
    add_paragraph_styled(doc, '功能描述：骑手接收待配送订单', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '操作流程：', bold=True)
    steps = [
        '在待接单列表中选择订单',
        '点击"接单"按钮',
        '系统确认接单成功',
        '订单状态变为待取货',
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '3.3 订单详情', level=2)
    add_paragraph_styled(doc, '功能描述：展示订单完整信息', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '信息结构：', bold=True)
    detail_items = [
        '订单基本信息：订单号、下单时间、订单状态',
        '收货信息：收货人、电话、详细地址',
        '商品明细：商品名称、数量',
        '配送信息：配送费、预计送达时间',
    ]
    for item in detail_items:
        add_list_item(doc, item)

    add_heading_styled(doc, '3.4 取货确认', level=2)
    add_paragraph_styled(doc, '功能描述：骑手到店后确认取货', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '操作流程：', bold=True)
    pickup_steps = [
        '到达商家门店',
        '点击"确认取货"按钮',
        '订单状态更新为待配送',
        '开始导航到收货地址',
    ]
    for i, step in enumerate(pickup_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    doc.add_page_break()

    add_heading_styled(doc, '3.5 开始配送', level=2)
    add_paragraph_styled(doc, '功能描述：更新订单状态为配送中', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '操作流程：', bold=True)
    delivery_steps = [
        '取货完成后自动进入配送状态',
        '或手动点击"开始配送"',
        '系统记录配送开始时间',
    ]
    for i, step in enumerate(delivery_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '3.6 送达确认', level=2)
    add_paragraph_styled(doc, '功能描述：确认订单已送达', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '操作流程：', bold=True)
    complete_steps = [
        '到达收货地址',
        '联系收货人确认',
        '点击"确认送达"按钮',
        '订单状态更新为已完成',
    ]
    for i, step in enumerate(complete_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '3.7 导航功能', level=2)
    add_paragraph_styled(doc, '功能描述：提供配送路线导航', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '功能特性：', bold=True)
    nav_items = [
        '显示从商家到收货地址的路线',
        '支持地图导航',
        '显示预计配送时间',
    ]
    for item in nav_items:
        add_list_item(doc, item)

    add_heading_styled(doc, '3.8 异常处理', level=2)
    add_paragraph_styled(doc, '功能描述：处理配送过程中的异常情况', indent=True)
    doc.add_paragraph()
    exp_headers = ['类型', '说明', '处理方式']
    exp_rows = [
        ['地址错误', '收货地址无法找到', '联系用户确认或联系商家'],
        ['联系不上', '无法联系收货人', '等待或联系商家'],
        ['商品问题', '商品损坏或缺失', '联系商家处理'],
    ]
    add_table_from_data(doc, exp_headers, exp_rows, first_col_bold=True)

    doc.add_page_break()

    # ===== 4. 个人中心模块 =====
    add_heading_styled(doc, '4. 个人中心模块', level=1)

    add_heading_styled(doc, '4.1 个人信息', level=2)
    add_paragraph_styled(doc, '功能描述：查看和修改个人资料', indent=True)
    doc.add_paragraph()
    info_field_headers = ['字段名', '类型', '说明']
    info_field_rows = [
        ['骑手姓名', '文本', '骑手真实姓名'],
        ['手机号', '文本', '联系电话'],
        ['身份证号', '文本', '实名认证信息'],
        ['头像', '图片', '个人头像'],
        ['配送等级', '文本', '骑手等级标识'],
    ]
    add_table_from_data(doc, info_field_headers, info_field_rows, first_col_bold=True)

    add_heading_styled(doc, '4.2 配送统计', level=2)
    add_paragraph_styled(doc, '功能描述：展示骑手配送业绩', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '统计内容：', bold=True)
    stats_items = [
        '今日配送单数',
        '本周配送单数',
        '本月配送单数',
        '累计配送单数',
        '完成率',
        '平均评分',
    ]
    for item in stats_items:
        add_list_item(doc, item)

    add_heading_styled(doc, '4.3 收入明细', level=2)
    add_paragraph_styled(doc, '功能描述：查看配送收入记录', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '收入结构：', bold=True)
    income_items = [
        '配送费收入',
        '奖励收入',
        '扣除项（如违规罚款）',
        '实发金额',
    ]
    for item in income_items:
        add_list_item(doc, item)

    add_heading_styled(doc, '4.4 消息通知', level=2)
    add_paragraph_styled(doc, '功能描述：接收系统通知', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, '消息类型：', bold=True)
    msg_items = [
        '订单消息：新订单提醒、订单取消通知',
        '系统消息：公告、活动通知',
        '个人消息：收入到账提醒',
    ]
    for item in msg_items:
        add_list_item(doc, item)

    doc.add_page_break()

    # ===== 5. 技术架构 =====
    add_heading_styled(doc, '5. 技术架构', level=1)

    add_heading_styled(doc, '5.1 技术栈', level=2)
    tech_headers = ['层级', '技术', '版本']
    tech_rows = [
        ['前端框架', 'React Native / 小程序', '-'],
        ['后端框架', 'Spring Boot', '2.x'],
        ['数据库', 'MySQL', '5.6+'],
        ['地图服务', '高德地图 / 百度地图', '-'],
    ]
    add_table_from_data(doc, tech_headers, tech_rows, first_col_bold=True)

    add_heading_styled(doc, '5.2 目录结构', level=2)
    add_code_block(doc, '''backend/
├── src/main/java/com/fr/
│   ├── controller/
│   │   └── RiderController.java    # 骑手端控制器
│   ├── service/
│   │   └── RiderService.java       # 骑手服务
│   ├── mapper/
│   │   └── RiderMapper.java        # 骑手数据访问
│   └── javaBean/
│       └── Rider.java              # 骑手实体类
└── src/main/resources/
    └── admin_docs/
        └── 骑手端功能需求文档.md''')

    doc.add_page_break()

    # ===== 6. 页面原型说明 =====
    add_heading_styled(doc, '6. 页面原型说明', level=1)

    add_heading_styled(doc, '6.1 骑手端首页布局', level=2)
    add_code_block(doc, '''┌─────────────────────────────────────┐
│         🍰 蛋糕骑手端               │
├─────────────────────────────────────┤
│  待接单 (3)    待配送 (2)    已完成  │
├─────────────────────────────────────┤
│  ┌───────────────────────────────┐ │
│  │ 订单号: D20260609001         │ │
│  │ 📦 3件商品                   │ │
│  │ 📍 XX街道XX小区XX栋          │ │
│  │ ⏱️ 预计30分钟送达            │ │
│  │ [接单]                       │ │
│  └───────────────────────────────┘ │
│  ┌───────────────────────────────┐ │
│  │ 订单号: D20260609002         │ │
│  │ 📦 2件商品                   │ │
│  │ 📍 YY路YY号                 │ │
│  │ ⏱️ 预计25分钟送达            │ │
│  │ [接单]                       │ │
│  └───────────────────────────────┘ │
├─────────────────────────────────────┤
│  🏠 首页    📋 订单    👤 我的    │
└─────────────────────────────────────┘''')

    add_heading_styled(doc, '6.2 页面清单', level=2)
    page_headers = ['页面路径', '页面名称', '功能说明']
    page_rows = [
        ['/rider/index', '首页', '订单列表（待接单、待配送、已完成）'],
        ['/rider/order/detail', '订单详情', '订单详细信息'],
        ['/rider/order/delivery', '配送导航', '配送路线导航'],
        ['/rider/profile', '个人中心', '个人信息和统计'],
        ['/rider/income', '收入明细', '收入记录'],
        ['/rider/messages', '消息通知', '系统消息'],
    ]
    add_table_from_data(doc, page_headers, page_rows, first_col_bold=True)

    doc.add_page_break()

    # ===== 附录 =====
    add_heading_styled(doc, '附录：订单状态流转', level=1)

    add_code_block(doc, '''待接单 → 待取货 → 待配送 → 已完成
            ↓(取消)
           已取消''')

    doc.add_paragraph()

    status_headers = ['状态码', '状态名称', '说明']
    status_rows = [
        ['1', '待接单', '订单已创建，等待骑手接单'],
        ['2', '待取货', '骑手已接单，等待取货'],
        ['3', '待配送', '已取货，正在配送中'],
        ['4', '已完成', '订单已送达'],
        ['0', '已取消', '订单已取消'],
    ]
    add_table_from_data(doc, status_headers, status_rows, first_col_bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— 蛋糕系统骑手端功能需求文档 · 版本 v1.0 —')
    set_chinese_font(run, font_name='宋体', size=10, color=(128, 128, 128))

    # 保存文档
    doc.save(output_path)
    print(f'骑手端需求文档已生成：{output_path}')


def generate_admin_requirements(output_path):
    """生成管理员需求文档Word文档"""
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('蛋糕系统')
    set_chinese_font(run, font_name='黑体', size=36, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('管理员功能需求文档')
    set_chinese_font(run, font_name='黑体', size=28, bold=True, color=(153, 0, 76))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        '文档版本：v1.0',
        '创建日期：2026年5月29日',
        '适用项目：CakeShop蛋糕系统',
    ]

    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_chinese_font(run, font_name='宋体', size=14)

    doc.add_page_break()

    # ===== 目录 =====
    add_heading_styled(doc, '目录', level=1)
    toc_items = [
        '1. 需求概述',
        '2. 功能需求清单',
        '3. 业务流程',
        '4. 数据权限',
        '5. 界面原型说明',
        '6. 非功能需求',
        '7. 与现有系统的集成',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_chinese_font(run, font_name='宋体', size=12)

    doc.add_page_break()

    # ===== 1. 需求概述 =====
    add_heading_styled(doc, '1. 需求概述', level=1)
    add_paragraph_styled(doc, '基于当前蛋糕系统已实现的用户端功能（商品浏览、购物车、订单管理等），管理员需要相应的后台管理功能来维护系统运营。', indent=True)

    doc.add_page_break()

    # ===== 2. 功能需求清单 =====
    add_heading_styled(doc, '2. 功能需求清单', level=1)

    add_heading_styled(doc, '2.1 管理员登录模块', level=2)
    login_headers = ['需求编号', '功能名称', '需求描述', '优先级']
    login_rows = [
        ['ADMIN-001', '管理员登录', '管理员通过账号密码登录后台管理系统', '高'],
        ['ADMIN-002', '安全退出', '管理员退出登录状态，清空session', '高'],
    ]
    add_table_from_data(doc, login_headers, login_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.2 用户管理模块', level=2)
    user_headers = ['需求编号', '功能名称', '需求描述', '优先级']
    user_rows = [
        ['ADMIN-003', '用户列表', '查看所有注册用户列表，支持分页', '高'],
        ['ADMIN-004', '用户详情', '查看单个用户的详细信息（用户名、电话、地址等）', '中'],
        ['ADMIN-005', '用户搜索', '根据用户名或电话搜索用户', '中'],
        ['ADMIN-006', '用户状态管理', '禁用/启用用户账号', '高'],
    ]
    add_table_from_data(doc, user_headers, user_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.3 商品管理模块', level=2)
    goods_headers = ['需求编号', '功能名称', '需求描述', '优先级']
    goods_rows = [
        ['ADMIN-007', '商品列表', '查看所有商品列表，支持分页和分类筛选', '高'],
        ['ADMIN-008', '添加商品', '添加新商品（名称、价格、库存、图片、分类等）', '高'],
        ['ADMIN-009', '编辑商品', '修改商品信息', '高'],
        ['ADMIN-010', '删除商品', '删除商品（需确认）', '高'],
        ['ADMIN-011', '商品搜索', '根据商品名称搜索', '中'],
        ['ADMIN-012', '库存管理', '批量调整商品库存', '高'],
    ]
    add_table_from_data(doc, goods_headers, goods_rows)

    doc.add_page_break()

    add_heading_styled(doc, '2.4 商品分类管理模块', level=2)
    type_headers = ['需求编号', '功能名称', '需求描述', '优先级']
    type_rows = [
        ['ADMIN-013', '分类列表', '查看所有商品分类', '高'],
        ['ADMIN-014', '添加分类', '添加新的商品分类', '高'],
        ['ADMIN-015', '编辑分类', '修改分类名称', '高'],
        ['ADMIN-016', '删除分类', '删除分类（需确保无商品关联）', '中'],
    ]
    add_table_from_data(doc, type_headers, type_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.5 订单管理模块', level=2)
    order_headers = ['需求编号', '功能名称', '需求描述', '优先级']
    order_rows = [
        ['ADMIN-017', '订单列表', '查看所有订单，支持分页和状态筛选', '高'],
        ['ADMIN-018', '订单详情', '查看订单详细信息（商品、收货人、金额等）', '高'],
        ['ADMIN-019', '订单搜索', '根据订单号或收货人搜索', '中'],
        ['ADMIN-020', '订单发货', '将订单状态从"已付款"改为"已发货"', '高'],
        ['ADMIN-021', '订单统计', '统计订单数量、总销售额等', '中'],
    ]
    add_table_from_data(doc, order_headers, order_rows)

    doc.add_paragraph()

    add_heading_styled(doc, '2.6 数据统计模块', level=2)
    stats_headers = ['需求编号', '功能名称', '需求描述', '优先级']
    stats_rows = [
        ['ADMIN-022', '销售统计', '统计指定时间范围内的销售额、订单数', '中'],
        ['ADMIN-023', '热销商品排行', '按销量排序的商品排行榜', '中'],
        ['ADMIN-024', '用户统计', '统计注册用户数量、活跃用户数', '低'],
    ]
    add_table_from_data(doc, stats_headers, stats_rows)

    doc.add_page_break()

    # ===== 3. 业务流程 =====
    add_heading_styled(doc, '3. 业务流程', level=1)

    add_heading_styled(doc, '3.1 订单处理流程', level=2)
    add_code_block(doc, '''用户下单 → 订单状态:已付款(2) → 管理员确认订单 → 管理员点击发货 → 订单状态:已发货(3) → 用户确认收货 → 订单状态:已完成(4)
                                                              ↓
                                                   用户取消订单 → 订单状态:已取消(5)''')

    doc.add_paragraph()

    add_heading_styled(doc, '3.2 商品管理流程', level=2)
    add_code_block(doc, '''添加商品 → 填写商品信息 → 上传图片 → 选择分类 → 保存 → 商品上架
                                                    ↓
                                              编辑/删除/调整库存''')

    doc.add_page_break()

    # ===== 4. 数据权限 =====
    add_heading_styled(doc, '4. 数据权限', level=1)

    perm_headers = ['角色', '权限范围']
    perm_rows = [
        ['超级管理员', '所有功能权限'],
        ['商品管理员', '商品管理、分类管理'],
        ['订单管理员', '订单管理、订单统计'],
        ['查看权限', '仅查看，无修改删除权限'],
    ]
    add_table_from_data(doc, perm_headers, perm_rows, first_col_bold=True)

    doc.add_page_break()

    # ===== 5. 界面原型说明 =====
    add_heading_styled(doc, '5. 界面原型说明', level=1)

    add_heading_styled(doc, '5.1 管理后台首页', level=2)
    layout_items = [
        '顶部：系统名称、登录用户、退出按钮',
        '左侧：功能菜单导航',
        '主内容区：数据统计概览（今日订单、今日销售额、商品总数、用户总数）',
    ]
    for item in layout_items:
        add_list_item(doc, item)

    add_heading_styled(doc, '5.2 列表页面通用布局', level=2)
    list_items = [
        '顶部：搜索框、筛选条件、新增按钮',
        '中部：数据表格（支持分页）',
        '底部：分页导航',
    ]
    for item in list_items:
        add_list_item(doc, item)

    add_heading_styled(doc, '5.3 表单页面通用布局', level=2)
    form_items = [
        '顶部：表单标题',
        '中部：表单字段',
        '底部：保存、取消按钮',
    ]
    for item in form_items:
        add_list_item(doc, item)

    doc.add_page_break()

    # ===== 6. 非功能需求 =====
    add_heading_styled(doc, '6. 非功能需求', level=1)

    nfr_headers = ['需求类型', '描述']
    nfr_rows = [
        ['安全性', '管理员密码加密存储，登录失败次数限制，操作日志记录'],
        ['响应速度', '页面加载时间不超过3秒'],
        ['兼容性', '支持主流浏览器（Chrome、Firefox、Safari）'],
        ['易用性', '操作流程清晰，提供必要的提示信息'],
    ]
    add_table_from_data(doc, nfr_headers, nfr_rows, first_col_bold=True)

    doc.add_page_break()

    # ===== 7. 与现有系统的集成 =====
    add_heading_styled(doc, '7. 与现有系统的集成', level=1)

    add_heading_styled(doc, '7.1 数据模型复用', level=2)
    model_headers = ['现有实体', '管理员功能使用场景']
    model_rows = [
        ['User', '用户管理模块'],
        ['Goods', '商品管理模块'],
        ['Type', '分类管理模块'],
        ['Order', '订单管理模块'],
        ['OrderItem', '订单详情展示'],
    ]
    add_table_from_data(doc, model_headers, model_rows, first_col_bold=True)

    doc.add_paragraph()

    add_heading_styled(doc, '7.2 状态码对应', level=2)
    status_code_headers = ['订单状态码', '状态名称', '管理员操作']
    status_code_rows = [
        ['2', '已付款', '可发货'],
        ['3', '已发货', '查看物流'],
        ['4', '已完成', '无操作'],
        ['5', '已取消', '无操作'],
    ]
    add_table_from_data(doc, status_code_headers, status_code_rows, first_col_bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— 蛋糕系统管理员功能需求文档 · 版本 v1.0 —')
    set_chinese_font(run, font_name='宋体', size=10, color=(128, 128, 128))

    # 保存文档
    doc.save(output_path)
    print(f'管理员需求文档已生成：{output_path}')


def generate_deployment_guide(output_path):
    """生成项目上线操作文档Word文档"""
    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('CakeShop 蛋糕系统')
    set_chinese_font(run, font_name='黑体', size=32, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('项目上线操作文档')
    set_chinese_font(run, font_name='黑体', size=28, bold=True, color=(0, 153, 76))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    info_lines = [
        '目标读者：Docker 和阿里云 0 基础用户',
        '操作难度：⭐⭐⭐（中等）',
        '预计时间：60-90分钟',
        '文档版本：v1.0',
        '创建时间：2026年6月23日',
    ]

    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        set_chinese_font(run, font_name='宋体', size=14)

    doc.add_page_break()

    # ===== 目录 =====
    add_heading_styled(doc, '目录', level=1)
    toc_items = [
        '一、本地 Docker 部署',
        '二、阿里云上线准备',
        '三、阿里云服务器配置',
        '四、项目部署到阿里云',
        '五、常见问题排查',
        '六、附录',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_chinese_font(run, font_name='宋体', size=12)

    doc.add_page_break()

    # ===== 一、本地 Docker 部署 =====
    add_heading_styled(doc, '一、本地 Docker 部署', level=1)

    add_heading_styled(doc, '1.1 环境准备', level=2)

    add_heading_styled(doc, '步骤1：安装 Docker Desktop', level=3)
    install_steps = [
        '下载 Docker Desktop：https://www.docker.com/products/docker-desktop/',
        '双击安装包进行安装',
        '安装完成后启动 Docker Desktop',
        '确保 Docker 图标显示在任务栏（Windows）或状态栏（Mac）',
    ]
    for i, step in enumerate(install_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '步骤2：配置镜像加速器（非常重要！）', level=3)
    add_paragraph_styled(doc, '由于国内访问 Docker Hub 速度较慢，需要配置镜像加速器。', indent=True)
    doc.add_paragraph()
    add_paragraph_styled(doc, 'Windows 系统配置方法：', bold=True)
    config_steps = [
        '右键点击任务栏中的 Docker 图标，选择「Settings」',
        '在左侧菜单中选择「Docker Engine」',
        '在配置框中添加镜像加速器配置',
        '点击「Apply & Restart」按钮重启 Docker',
    ]
    for i, step in enumerate(config_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    doc.add_paragraph()
    add_paragraph_styled(doc, '验证配置是否生效：', bold=True)
    add_code_block(doc, 'docker info')
    add_paragraph_styled(doc, '如果看到 Registry Mirrors 列表包含你配置的镜像地址，说明配置成功。', indent=True)

    add_heading_styled(doc, '步骤3：验证 Docker 安装', level=3)
    add_paragraph_styled(doc, '打开命令提示符（Windows）或终端（Mac），执行以下命令：', indent=True)
    add_code_block(doc, '''docker --version
docker-compose --version''')
    add_paragraph_styled(doc, '如果显示版本信息，说明安装成功。', indent=True)

    doc.add_page_break()

    add_heading_styled(doc, '1.2 创建 Docker 配置文件', level=2)
    add_paragraph_styled(doc, '在项目根目录下创建以下文件：', indent=True)

    add_heading_styled(doc, '文件1：Dockerfile', level=3)
    add_code_block(doc, '''# 使用多阶段构建
FROM maven:3.8.6-openjdk-17 AS builder

# 设置工作目录
WORKDIR /app

# 复制 Maven 配置文件和源代码
COPY pom.xml .
COPY src ./src

# 编译项目，跳过测试
RUN mvn clean package -DskipTests

# 第二阶段：运行环境
FROM openjdk:17-jdk-slim

# 设置工作目录
WORKDIR /app

# 从构建阶段复制编译好的 JAR 文件
COPY --from=builder /app/target/*.jar app.jar

# 暴露端口
EXPOSE 8090

# 设置时区
ENV TZ=Asia/Shanghai

# 启动命令
ENTRYPOINT ["java", "-jar", "app.jar"]''')

    doc.add_page_break()

    add_heading_styled(doc, '文件2：docker-compose.yml', level=3)
    add_code_block(doc, '''version: '3.8'

services:
  # MySQL 数据库服务
  mysql:
    image: mysql:8.0
    container_name: cakeshop-mysql
    restart: always
    environment:
      MYSQL_ROOT_PASSWORD: your_mysql_password  # 修改为你的密码
      MYSQL_DATABASE: cookieshop
      MYSQL_USER: cakeshop_user
      MYSQL_PASSWORD: cakeshop_password
    ports:
      - "3306:3306"
    volumes:
      - mysql_data:/var/lib/mysql
    networks:
      - cakeshop-network

  # CakeShop 应用服务
  cakeshop:
    build: .
    container_name: cakeshop-app
    restart: always
    environment:
      SPRING_DATASOURCE_URL: jdbc:mysql://mysql:3306/cookieshop?characterEncoding=UTF-8&serverTimezone=UTC
      SPRING_DATASOURCE_USERNAME: cakeshop_user
      SPRING_DATASOURCE_PASSWORD: cakeshop_password
    ports:
      - "8090:8090"
    depends_on:
      - mysql
    networks:
      - cakeshop-network

volumes:
  mysql_data:

networks:
  cakeshop-network:
    driver: bridge''')

    doc.add_page_break()

    add_heading_styled(doc, '文件3：.dockerignore（可选）', level=3)
    add_code_block(doc, '''target/
.mvn/
.git/
*.log''')

    add_heading_styled(doc, '1.3 构建并启动容器', level=2)

    add_heading_styled(doc, '步骤1：打开命令行', level=3)
    add_paragraph_styled(doc, '在 Windows 上打开命令提示符或 PowerShell，切换到项目目录：', indent=True)
    add_code_block(doc, 'cd c:\\1\\CakeShop')

    add_heading_styled(doc, '步骤2：启动容器', level=3)
    add_paragraph_styled(doc, '执行以下命令启动所有服务：', indent=True)
    add_code_block(doc, 'docker-compose up -d')
    add_paragraph_styled(doc, '说明：-d 参数表示后台运行，第一次运行会下载镜像，可能需要几分钟。', indent=True)

    add_heading_styled(doc, '步骤3：查看容器状态', level=3)
    add_code_block(doc, 'docker-compose ps')
    add_paragraph_styled(doc, '如果显示 cakeshop-mysql 和 cakeshop-app 都是 Up 状态，说明启动成功。', indent=True)

    add_heading_styled(doc, '1.4 验证本地部署', level=2)

    add_heading_styled(doc, '步骤1：等待服务启动', level=3)
    add_paragraph_styled(doc, '容器启动后，应用需要一些时间初始化。等待约 30 秒后进行验证。', indent=True)

    add_heading_styled(doc, '步骤2：访问应用', level=3)
    add_paragraph_styled(doc, '打开浏览器，访问：', indent=True)
    add_code_block(doc, 'http://localhost:8090')
    add_paragraph_styled(doc, '如果能看到蛋糕系统首页，说明本地部署成功！', indent=True)

    add_heading_styled(doc, '步骤3：测试功能', level=3)
    test_items = [
        '点击导航栏的各个链接',
        '尝试注册新用户',
        '浏览商品列表',
        '添加商品到购物车',
    ]
    for item in test_items:
        add_list_item(doc, item)

    doc.add_page_break()

    # ===== 二、阿里云上线准备 =====
    add_heading_styled(doc, '二、阿里云上线准备', level=1)

    add_heading_styled(doc, '2.1 注册阿里云账号', level=2)
    reg_steps = [
        '打开阿里云官网：https://www.aliyun.com/',
        '点击右上角「免费注册」',
        '按照提示完成账号注册和实名认证',
    ]
    for i, step in enumerate(reg_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '2.2 申请免费试用 ECS', level=2)

    add_heading_styled(doc, '步骤1：进入 ECS 控制台', level=3)
    ecs_steps = [
        '登录阿里云控制台：https://ecs.console.aliyun.com/',
        '如果没有开通 ECS，会提示开通，点击「立即开通」',
    ]
    for i, step in enumerate(ecs_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '步骤2：申请免费试用', level=3)
    trial_steps = [
        '在 ECS 控制台首页，找到「免费试用」入口',
        '选择「云服务器 ECS」免费试用',
        '配置信息：地域选择离你最近的区域，选择免费规格，操作系统选择 Ubuntu 22.04 LTS 或 CentOS 7.x，购买时长选择 1 个月',
    ]
    for i, step in enumerate(trial_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '步骤3：设置登录密码', level=3)
    add_paragraph_styled(doc, '非常重要！在创建实例时设置登录密码：', indent=True)
    pwd_steps = [
        '在「实例配置」页面找到「登录凭证」',
        '选择「自定义密码」',
        '设置一个安全的密码（建议包含大小写字母、数字、特殊字符）',
        '记住这个密码，后续登录服务器需要使用',
    ]
    for i, step in enumerate(pwd_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '2.3 配置安全组', level=2)

    add_heading_styled(doc, '步骤1：找到安全组配置', level=3)
    sg_steps = [
        '在 ECS 控制台左侧菜单，点击「安全组」',
        '选择你创建的实例所在的安全组',
    ]
    for i, step in enumerate(sg_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '步骤2：添加安全规则', level=3)
    add_paragraph_styled(doc, '点击「添加安全规则」，添加以下规则：', indent=True)
    sg_headers = ['规则方向', '端口范围', '授权对象', '说明']
    sg_rows = [
        ['入方向', '22', '0.0.0.0/0', 'SSH 登录'],
        ['入方向', '80', '0.0.0.0/0', 'HTTP 访问'],
        ['入方向', '443', '0.0.0.0/0', 'HTTPS 访问'],
        ['入方向', '8090', '0.0.0.0/0', '应用端口'],
        ['入方向', '3306', '0.0.0.0/0', 'MySQL 端口（生产环境建议限制 IP）'],
    ]
    add_table_from_data(doc, sg_headers, sg_rows)

    doc.add_page_break()

    # ===== 三、阿里云服务器配置 =====
    add_heading_styled(doc, '三、阿里云服务器配置', level=1)

    add_heading_styled(doc, '3.1 登录 ECS 服务器', level=2)

    add_heading_styled(doc, '方法1：使用阿里云控制台登录', level=3)
    console_login_steps = [
        '在 ECS 控制台，找到你的实例',
        '点击「远程连接」',
        '选择「Workbench 远程连接」',
        '输入登录密码，点击「确定」',
    ]
    for i, step in enumerate(console_login_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '方法2：使用 SSH 工具登录（推荐）', level=3)
    add_paragraph_styled(doc, 'Windows 用户：使用 PuTTY 或 Windows Terminal', indent=True)
    add_paragraph_styled(doc, 'Mac/Linux 用户：使用终端', indent=True)
    add_code_block(doc, 'ssh root@你的服务器公网IP')
    add_paragraph_styled(doc, '输入之前设置的密码即可登录。', indent=True)

    add_heading_styled(doc, '3.2 安装 Docker', level=2)

    add_heading_styled(doc, '步骤1：更新系统', level=3)
    add_code_block(doc, '''# Ubuntu 系统
apt update && apt upgrade -y

# CentOS 系统
yum update -y''')

    add_heading_styled(doc, '步骤2：安装 Docker（Ubuntu）', level=3)
    add_code_block(doc, '''# 安装依赖
apt install -y apt-transport-https ca-certificates curl software-properties-common

# 添加 Docker GPG 密钥
curl -fsSL https://download.docker.com/linux/ubuntu/gpg | gpg --dearmor -o /usr/share/keyrings/docker-archive-keyring.gpg

# 添加 Docker 源
echo "deb [arch=$(dpkg --print-architecture) signed-by=/usr/share/keyrings/docker-archive-keyring.gpg] https://download.docker.com/linux/ubuntu $(lsb_release -cs) stable" | tee /etc/apt/sources.list.d/docker.list > /dev/null

# 安装 Docker
apt update && apt install -y docker-ce docker-ce-cli containerd.io

# 启动 Docker 服务
systemctl start docker
systemctl enable docker

# 验证安装
docker --version''')

    doc.add_page_break()

    add_heading_styled(doc, '步骤2：安装 Docker（CentOS）', level=3)
    add_code_block(doc, '''# 安装依赖
yum install -y yum-utils device-mapper-persistent-data lvm2

# 添加 Docker 源
yum-config-manager --add-repo https://download.docker.com/linux/centos/docker-ce.repo

# 安装 Docker
yum install -y docker-ce docker-ce-cli containerd.io

# 启动 Docker 服务
systemctl start docker
systemctl enable docker

# 验证安装
docker --version''')

    add_heading_styled(doc, '步骤3：安装 Docker Compose', level=3)
    add_code_block(doc, '''# 下载 Docker Compose
curl -SL https://github.com/docker/compose/releases/download/v2.20.0/docker-compose-linux-x86_64 -o /usr/local/bin/docker-compose

# 添加执行权限
chmod +x /usr/local/bin/docker-compose

# 验证安装
docker-compose --version''')

    add_heading_styled(doc, '3.3 安装 MySQL（可选，也可以用 Docker）', level=2)
    add_paragraph_styled(doc, '方案A：使用 Docker MySQL（推荐）- 在后续步骤中通过 docker-compose 自动安装。', indent=True)
    add_paragraph_styled(doc, '方案B：安装本地 MySQL：', indent=True)
    add_code_block(doc, '''# Ubuntu
apt install -y mysql-server

# CentOS
yum install -y mysql-server

# 启动服务
systemctl start mysql
systemctl enable mysql''')

    doc.add_page_break()

    # ===== 四、项目部署到阿里云 =====
    add_heading_styled(doc, '四、项目部署到阿里云', level=1)

    add_heading_styled(doc, '4.1 上传项目代码', level=2)

    add_heading_styled(doc, '方法1：使用 Git 克隆', level=3)
    add_code_block(doc, '''# 创建项目目录
mkdir -p /opt/cakeshop
cd /opt/cakeshop

# 克隆代码（如果有 Git 仓库）
git clone <你的仓库地址> .''')

    add_heading_styled(doc, '方法2：使用 FTP 上传', level=3)
    ftp_steps = [
        '在本地打包项目',
        '使用 FileZilla 或 WinSCP 上传压缩包到服务器',
        '在服务器解压',
    ]
    for i, step in enumerate(ftp_steps, 1):
        p = doc.add_paragraph(style='List Number')
        run = p.add_run(step)
        set_chinese_font(run, font_name='宋体', size=12)

    add_heading_styled(doc, '4.2 配置环境变量', level=2)
    add_paragraph_styled(doc, '编辑 docker-compose.yml，修改数据库密码：', indent=True)
    add_code_block(doc, 'nano docker-compose.yml')
    add_paragraph_styled(doc, '修改以下内容：', indent=True)
    config_items = [
        'MYSQL_ROOT_PASSWORD：设置一个安全的密码',
        'MYSQL_PASSWORD：设置数据库用户密码',
        'SPRING_DATASOURCE_PASSWORD：与上面一致',
    ]
    for item in config_items:
        add_list_item(doc, item)
    add_paragraph_styled(doc, '按 Ctrl+O 保存，按 Ctrl+X 退出。', indent=True)

    add_heading_styled(doc, '4.3 构建并启动容器', level=2)
    add_code_block(doc, '''# 进入项目目录
cd /opt/cakeshop

# 构建并启动容器
docker-compose up -d''')
    add_paragraph_styled(doc, '注意：第一次运行需要下载镜像和构建，可能需要 5-10 分钟。', indent=True)

    add_heading_styled(doc, '4.4 验证上线', level=2)

    add_heading_styled(doc, '步骤1：查看容器状态', level=3)
    add_code_block(doc, 'docker-compose ps')

    add_heading_styled(doc, '步骤2：查看应用日志', level=3)
    add_code_block(doc, 'docker-compose logs cakeshop')
    add_paragraph_styled(doc, '如果看到 Started CakeShopApplication，说明启动成功。', indent=True)

    add_heading_styled(doc, '步骤3：访问应用', level=3)
    add_paragraph_styled(doc, '打开浏览器，访问：', indent=True)
    add_code_block(doc, 'http://你的服务器公网IP:8090')
    add_paragraph_styled(doc, '如果能看到蛋糕系统首页，说明上线成功！', indent=True)

    doc.add_page_break()

    # ===== 五、常见问题排查 =====
    add_heading_styled(doc, '五、常见问题排查', level=1)

    add_heading_styled(doc, '问题1：Docker 容器启动失败', level=2)
    add_paragraph_styled(doc, '现象：docker-compose ps 显示容器状态为 Exit', indent=True)
    add_paragraph_styled(doc, '解决方法：', indent=True)
    add_code_block(doc, '''# 查看容器日志
docker-compose logs cakeshop

# 常见原因：
# 1. 数据库连接失败 - 检查数据库配置
# 2. 端口被占用 - 检查 8090 端口
# 3. 代码编译错误 - 检查 Maven 编译日志''')

    add_heading_styled(doc, '问题2：无法访问应用', level=2)
    add_paragraph_styled(doc, '现象：浏览器显示无法访问', indent=True)
    add_paragraph_styled(doc, '解决方法：', indent=True)
    access_items = [
        '检查安全组是否开放了 8090 端口',
        '检查服务器防火墙是否开放了端口',
    ]
    for item in access_items:
        add_list_item(doc, item)
    add_code_block(doc, '''# Ubuntu
ufw status
ufw allow 8090

# CentOS
firewall-cmd --list-all
firewall-cmd --add-port=8090/tcp --permanent
firewall-cmd --reload''')

    add_heading_styled(doc, '问题3：数据库连接失败', level=2)
    add_paragraph_styled(doc, '现象：应用日志显示 Cannot get connection', indent=True)
    add_paragraph_styled(doc, '解决方法：', indent=True)
    db_items = [
        '检查 docker-compose.yml 中的数据库配置',
        '确保数据库容器先启动',
    ]
    for item in db_items:
        add_list_item(doc, item)
    add_code_block(doc, '''docker-compose up -d mysql
# 等待 10 秒后启动应用
docker-compose up -d cakeshop''')

    add_heading_styled(doc, '问题4：Docker 命令权限不足', level=2)
    add_paragraph_styled(doc, '现象：docker: Got permission denied', indent=True)
    add_paragraph_styled(doc, '解决方法：', indent=True)
    add_code_block(doc, '''# 将当前用户添加到 docker 组
sudo usermod -aG docker $USER

# 重新登录或执行
newgrp docker''')

    doc.add_page_break()

    # ===== 六、附录 =====
    add_heading_styled(doc, '六、附录', level=1)

    add_heading_styled(doc, 'A. 常用命令速查表', level=2)
    add_code_block(doc, '''# Docker 相关
docker ps                    # 查看运行中的容器
docker ps -a                 # 查看所有容器
docker stop <容器名>          # 停止容器
docker start <容器名>         # 启动容器
docker rm <容器名>            # 删除容器
docker logs <容器名>          # 查看容器日志
docker-compose up -d         # 启动所有服务
docker-compose down          # 停止并删除容器
docker-compose restart       # 重启服务

# 服务器相关
systemctl status docker      # 查看 Docker 状态
systemctl restart docker     # 重启 Docker 服务
netstat -tlnp                # 查看端口占用''')

    add_heading_styled(doc, 'B. 文件结构说明', level=2)
    add_code_block(doc, '''CakeShop/
├── Dockerfile              # Docker 构建配置
├── docker-compose.yml      # 容器编排配置
├── .dockerignore           # Docker 忽略文件
├── pom.xml                 # Maven 依赖配置
├── src/
│   └── main/
│       ├── java/           # Java 源代码
│       └── resources/
│           ├── application.yml  # 应用配置
│           ├── static/     # 静态资源
│           └── templates/  # 页面模板
└── target/                 # 编译产物（运行时生成）''')

    add_heading_styled(doc, 'C. 端口说明', level=2)
    port_headers = ['端口', '服务', '说明']
    port_rows = [
        ['80', 'HTTP', '网站访问（可配置 Nginx 反向代理）'],
        ['443', 'HTTPS', '加密访问（需要 SSL 证书）'],
        ['8090', 'CakeShop', '应用默认端口'],
        ['3306', 'MySQL', '数据库端口'],
        ['22', 'SSH', '远程登录'],
    ]
    add_table_from_data(doc, port_headers, port_rows, first_col_bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('— CakeShop 蛋糕系统项目上线操作文档 · 版本 v1.0 —')
    set_chinese_font(run, font_name='宋体', size=10, color=(128, 128, 128))

    # 保存文档
    doc.save(output_path)
    print(f'项目上线操作文档已生成：{output_path}')


if __name__ == '__main__':
    word_dir = r'c:\24级软件技术3班_2400130315_黄俊鹏\CakeShop\docs\word'

    automation_test_path = f'{word_dir}\\自动化测试报告.docx'
    rider_req_path = f'{word_dir}\\骑手端需求文档.docx'
    admin_req_path = f'{word_dir}\\管理员需求文档.docx'
    deployment_path = f'{word_dir}\\项目上线操作文档.docx'

    generate_automation_test_report(automation_test_path)
    generate_rider_requirements(rider_req_path)
    generate_admin_requirements(admin_req_path)
    generate_deployment_guide(deployment_path)

    print('\n✅ 所有Word文档生成完成！')
    print(f'文档存放位置：{word_dir}')
